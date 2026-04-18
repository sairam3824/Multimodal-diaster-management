from __future__ import annotations

import json
import shutil
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from thesis_content import ABSTRACT, ACKNOWLEDGEMENT, ACRONYMS, APPENDICES, CHAPTERS, REFERENCES


ROOT = Path(__file__).resolve().parent.parent
THESIS_DIR = ROOT / "thesis"
METADATA_PATH = THESIS_DIR / "metadata.json"
OUTPUT_PATH = THESIS_DIR / "Multimodal_Disaster_Intelligence_Thesis_VIT.docx"
FIGURES_DIR = THESIS_DIR / "figures"


def load_metadata() -> dict:
    return json.loads(METADATA_PATH.read_text(encoding="utf-8"))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Body Text", "List Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading3.font.size = Pt(12)
    heading3.font.bold = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(1.0)


def add_body_paragraph(doc: Document, text: str):
    para = doc.add_paragraph(text)
    style_body(para)
    return para


def add_center_block(doc: Document, lines: list[tuple[str, int, bool, bool]], space_after=6):
    for text, size, bold, italic in lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def add_heading(doc: Document, text: str, level: int):
    para = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_field(paragraph, instruction: str, placeholder: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def add_sequence(run, label: str):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"SEQ {label} \\* ARABIC"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def copy_figure(source_rel_path: str) -> Path:
    source = ROOT / source_rel_path
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    target_name = source_rel_path.replace("/", "__").replace(" ", "_")
    target = FIGURES_DIR / target_name
    shutil.copy2(source, target)
    return target


def add_figure(doc: Document, source_rel_path: str, caption: str, width_inches: float):
    figure_path = copy_figure(source_rel_path)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(figure_path), width=Inches(width_inches))
    caption_para = doc.add_paragraph(style="Caption")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run("Figure ")
    add_sequence(run, "Figure")
    caption_para.add_run(f". {caption}")


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]):
    caption_para = doc.add_paragraph(style="Caption")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run("Table ")
    add_sequence(run, "Table")
    clean_title = re.sub(r"^Table\s+\d+(?:\.\d+)?\s*", "", title, flags=re.IGNORECASE)
    caption_para.add_run(f". {clean_title}")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.15
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_cover_page(doc: Document, metadata: dict):
    add_center_block(
        doc,
        [
            ("A Project Report on", 14, False, True),
            ("", 12, False, False),
            (metadata["title"].upper(), 18, True, False),
            ("", 12, False, False),
            ("Submitted in partial fulfillment of the requirements for the award of the degree of", 12, False, False),
            (f"{metadata['degree']} in {metadata['programme']}", 16, True, False),
            ("", 12, False, False),
            ("by", 12, False, False),
        ],
        space_after=8,
    )

    for candidate in metadata["candidates"]:
        add_center_block(
            doc,
            [(f"{candidate['name']} ({candidate['registration_number']})", 13, True, False)],
            space_after=4,
        )

    add_center_block(
        doc,
        [
            ("", 12, False, False),
            (metadata["school"], 14, True, False),
            (metadata["institution"], 14, True, False),
            (metadata["campus"], 12, False, False),
            (metadata["submission_month_year"], 13, True, False),
        ],
        space_after=6,
    )
    doc.add_page_break()


def add_title_page(doc: Document, metadata: dict):
    add_center_block(
        doc,
        [
            (metadata["title"].upper(), 18, True, False),
            ("", 12, False, False),
            ("Submitted in partial fulfillment of the requirements for the award of the degree of", 12, False, False),
            (f"{metadata['degree']} in {metadata['programme']}", 16, True, False),
            ("", 12, False, False),
            ("by", 12, False, False),
        ],
        space_after=8,
    )
    for candidate in metadata["candidates"]:
        add_center_block(doc, [(f"{candidate['name']} ({candidate['registration_number']})", 13, True, False)], space_after=4)

    add_center_block(
        doc,
        [
            ("", 12, False, False),
            (metadata["school"], 14, True, False),
            (metadata["institution"], 14, True, False),
            (metadata["campus"], 12, False, False),
            (metadata["submission_month_year"], 13, True, False),
        ],
        space_after=6,
    )
    doc.add_page_break()


def add_declaration(doc: Document, metadata: dict):
    add_heading(doc, "DECLARATION", 1)
    candidate_names = ", ".join(
        f"{c['name']} ({c['registration_number']})" for c in metadata["candidates"]
    )
    text = (
        f"We hereby declare that the thesis entitled \"{metadata['title']}\" submitted by "
        f"{candidate_names} for the award of the degree of {metadata['degree']} in {metadata['programme']} "
        f"at {metadata['institution']} is a record of bonafide work carried out under the guidance of "
        f"{metadata['guide_name']}. The work reported in this thesis has not been submitted, either in full "
        f"or in part, for the award of any other degree or diploma in this institute or any other institution."
    )
    add_body_paragraph(doc, text)
    doc.add_paragraph()
    add_body_paragraph(doc, f"Place: {metadata['place']}")
    add_body_paragraph(doc, "Date: ____________________")
    add_body_paragraph(doc, "Signature(s) of the Candidate(s): ______________________________")
    doc.add_page_break()


def add_certificate(doc: Document, metadata: dict):
    add_heading(doc, "CERTIFICATE", 1)
    candidate_names = ", ".join(
        f"{c['name']} ({c['registration_number']})" for c in metadata["candidates"]
    )
    text = (
        f"This is to certify that the senior design project titled \"{metadata['title']}\" submitted by "
        f"{candidate_names} in partial fulfillment of the requirements for the award of {metadata['degree']} "
        f"in {metadata['programme']} at {metadata['institution']} is a record of bonafide work carried out "
        f"under my guidance. To the best of my knowledge, the contents of this thesis, in full or in part, "
        f"have neither been copied from any other source nor submitted to any other institute or university "
        f"for the award of any degree or diploma."
    )
    add_body_paragraph(doc, text)
    doc.add_paragraph()
    add_body_paragraph(doc, f"{metadata['guide_name']}")
    add_body_paragraph(doc, f"{metadata['guide_designation']}")
    add_body_paragraph(doc, "Internal Guide")
    doc.add_paragraph()
    add_body_paragraph(doc, "Internal Examiner: ____________________")
    add_body_paragraph(doc, "External Examiner: ____________________")
    doc.add_paragraph()
    add_body_paragraph(doc, f"Approved by: {metadata['dean_name']}, Dean, {metadata['school']}")
    doc.add_page_break()


def add_abstract(doc: Document):
    add_heading(doc, "ABSTRACT", 1)
    for paragraph in ABSTRACT.split("\n\n"):
        add_body_paragraph(doc, paragraph.strip())
    doc.add_page_break()


def add_acknowledgement(doc: Document, metadata: dict):
    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    text = ACKNOWLEDGEMENT.format(
        guide_name=metadata["guide_name"],
        guide_designation=metadata["guide_designation"],
        institution=metadata["institution"],
        school=metadata["school"],
    )
    for paragraph in text.split("\n\n"):
        add_body_paragraph(doc, paragraph.strip())
    add_body_paragraph(doc, f"Place: {metadata['place']}")
    add_body_paragraph(doc, "Date: ____________________")
    doc.add_page_break()


def add_contents_pages(doc: Document):
    add_heading(doc, "TABLE OF CONTENTS", 1)
    para = doc.add_paragraph()
    add_field(para, 'TOC \\o "1-3" \\h \\z \\u', "Update field in Word to generate the table of contents.")
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", 1)
    para = doc.add_paragraph()
    add_field(para, 'TOC \\h \\z \\c "Figure"', "Update field in Word to generate the list of figures.")
    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES", 1)
    para = doc.add_paragraph()
    add_field(para, 'TOC \\h \\z \\c "Table"', "Update field in Word to generate the list of tables.")
    doc.add_page_break()


def add_acronyms(doc: Document):
    add_heading(doc, "LIST OF ACRONYMS", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Acronym"
    table.rows[0].cells[1].text = "Meaning"
    for acronym, meaning in ACRONYMS:
        row = table.add_row().cells
        row[0].text = acronym
        row[1].text = meaning
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.15
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
    doc.add_page_break()


def add_chapters(doc: Document):
    for chapter in CHAPTERS:
        add_heading(doc, chapter["title"], 1)
        for element in chapter["elements"]:
            if element["type"] == "section":
                add_heading(doc, element["title"].upper(), 2)
                for paragraph in element["paragraphs"]:
                    add_body_paragraph(doc, paragraph)
            elif element["type"] == "figure":
                add_figure(doc, element["path"], element["caption"], element["width_inches"])
            elif element["type"] == "table":
                add_table(doc, element["title"], element["headers"], element["rows"])
        doc.add_page_break()


def add_references(doc: Document):
    add_heading(doc, "REFERENCES", 1)
    for item in REFERENCES:
        para = doc.add_paragraph(item)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.left_indent = Cm(0.75)
        para.paragraph_format.first_line_indent = Cm(-0.75)
        para.paragraph_format.space_after = Pt(4)
    doc.add_page_break()


def add_appendices(doc: Document):
    add_heading(doc, "APPENDICES", 1)
    for appendix in APPENDICES:
        add_heading(doc, appendix["title"], 2)
        for paragraph in appendix["paragraphs"]:
            add_body_paragraph(doc, paragraph)


def add_footer_page_numbers(doc: Document):
    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        add_page_number(footer_para)


def main():
    metadata = load_metadata()
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    for existing in FIGURES_DIR.iterdir():
        if existing.is_file():
            existing.unlink()
    doc = Document()
    configure_document(doc)
    set_update_fields_on_open(doc)
    add_cover_page(doc, metadata)
    add_title_page(doc, metadata)
    add_declaration(doc, metadata)
    add_certificate(doc, metadata)
    add_abstract(doc)
    add_acknowledgement(doc, metadata)
    add_contents_pages(doc)
    add_acronyms(doc)
    add_chapters(doc)
    add_references(doc)
    add_appendices(doc)
    add_footer_page_numbers(doc)
    doc.save(OUTPUT_PATH)
    print(f"Generated thesis at: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
