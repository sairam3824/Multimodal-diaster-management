#!/usr/bin/env python3
"""
Generate VIT-AP Thesis Format Word Document for
Tri-Modal Fusion Disaster Intelligence Project
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# PAGE SETUP: A4 with VIT-AP margins
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

# ============================================================
# STYLE DEFINITIONS
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def set_font(run, size=12, bold=False, italic=False, underline=False, caps=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if caps:
        run.font.all_caps = True

def add_centered_para(text, size=12, bold=False, italic=False, underline=False, caps=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, underline=underline, caps=caps)
    return p

def add_normal_para(text, size=12, bold=False, italic=False, space_after=6, space_before=0, first_line_indent=None, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading_custom(text, level=1):
    """Chapter=16pt bold, Section=14pt CAPS, Subsection=12pt CAPS"""
    p = doc.add_paragraph()
    if level == 0:  # Chapter heading
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=16, bold=True, caps=True)
    elif level == 1:  # Section heading
        run = p.add_run(text.upper())
        set_font(run, size=14, bold=True)
    elif level == 2:  # Subsection heading
        run = p.add_run(text.upper())
        set_font(run, size=12, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_page_break():
    doc.add_page_break()

def add_empty_lines(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def add_table_row(table, cells_data, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = alignment
        run = p.add_run(str(text))
        set_font(run, size=11, bold=bold)
    return row

# ============================================================
# COVER PAGE
# ============================================================
add_empty_lines(3)
add_centered_para('A project report on', size=12, italic=True, space_after=12)
add_empty_lines(1)
add_centered_para(
    'TRI-MODAL FUSION FOR REAL-TIME DISASTER INTELLIGENCE:\n'
    'INTEGRATING ENVIRONMENTAL SENSOR METADATA, SATELLITE\n'
    'IMAGERY, AND SOCIAL MEDIA THROUGH ADAPTIVE CROSS-ATTENTION',
    size=20, bold=True, space_after=18
)
add_empty_lines(1)
add_centered_para('Submitted in partial fulfillment for the award of the degree of', size=14, italic=True, space_after=12)
add_centered_para('Bachelor of Technology', size=22, bold=True, space_after=18)
add_empty_lines(1)
add_centered_para('by', size=14, italic=True, space_after=12)
add_centered_para('MARURI SAI RAMA LINGA REDDY (Reg. No. XXXXXXXXX)', size=16, bold=True, space_after=24)
add_empty_lines(2)
add_centered_para('School of Computer Science and Engineering', size=16, bold=True, space_after=6)
add_centered_para('VIT-AP University', size=16, bold=True, space_after=6)
add_centered_para('Amaravati, Andhra Pradesh', size=12, space_after=12)
add_centered_para('April, 2026', size=12)

# ============================================================
# TITLE PAGE (same as cover but with guide info)
# ============================================================
add_page_break()
add_empty_lines(3)
add_centered_para('A project report on', size=12, italic=True, space_after=12)
add_empty_lines(1)
add_centered_para(
    'TRI-MODAL FUSION FOR REAL-TIME DISASTER INTELLIGENCE:\n'
    'INTEGRATING ENVIRONMENTAL SENSOR METADATA, SATELLITE\n'
    'IMAGERY, AND SOCIAL MEDIA THROUGH ADAPTIVE CROSS-ATTENTION',
    size=20, bold=True, space_after=18
)
add_empty_lines(1)
add_centered_para('Submitted in partial fulfillment for the award of the degree of', size=14, italic=True, space_after=12)
add_centered_para('Bachelor of Technology', size=22, bold=True, space_after=18)
add_empty_lines(1)
add_centered_para('by', size=14, italic=True, space_after=12)
add_centered_para('MARURI SAI RAMA LINGA REDDY (Reg. No. XXXXXXXXX)', size=16, bold=True, space_after=12)
add_empty_lines(1)
add_centered_para('Under the guidance of', size=12, italic=True, space_after=6)
add_centered_para('Dr. [GUIDE NAME]', size=14, bold=True, space_after=24)
add_empty_lines(1)
add_centered_para('School of Computer Science and Engineering', size=16, bold=True, space_after=6)
add_centered_para('VIT-AP University', size=16, bold=True, space_after=6)
add_centered_para('Amaravati, Andhra Pradesh', size=12, space_after=12)
add_centered_para('April, 2026', size=12)

# ============================================================
# DECLARATION
# ============================================================
add_page_break()
add_empty_lines(2)
add_centered_para('DECLARATION', size=14, bold=True, underline=True, space_after=24)
add_empty_lines(1)

add_normal_para(
    'I hereby declare that the thesis entitled "Tri-Modal Fusion for Real-Time Disaster Intelligence: '
    'Integrating Environmental Sensor Metadata, Satellite Imagery, and Social Media Through Adaptive '
    'Cross-Attention" submitted by me, for the award of the degree of Bachelor of Technology in Computer '
    'Science and Engineering to VIT-AP University is a record of bonafide work carried out by me under '
    'the supervision of Dr. [Guide Name].',
    space_after=12, first_line_indent=1.27
)

add_normal_para(
    'I further declare that the work reported in this thesis has not been submitted and will not be '
    'submitted, either in part or in full, for the award of any other degree or diploma in this '
    'institute or any other institute or university.',
    space_after=24, first_line_indent=1.27
)

add_empty_lines(4)
add_normal_para('Place: Amaravati', bold=True)
add_normal_para('Date:', bold=True, space_after=0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Signature of the Candidate')
set_font(run, size=12, bold=True)

# ============================================================
# CERTIFICATE
# ============================================================
add_page_break()
add_empty_lines(2)
add_centered_para('CERTIFICATE', size=14, bold=True, underline=True, space_after=24)
add_empty_lines(1)

add_normal_para(
    'This is to certify that the Senior Design Project titled "Tri-Modal Fusion for Real-Time '
    'Disaster Intelligence: Integrating Environmental Sensor Metadata, Satellite Imagery, and Social '
    'Media Through Adaptive Cross-Attention" that is being submitted by MARURI SAI RAMA LINGA REDDY '
    '(Reg. No. XXXXXXXXX) is in partial fulfillment of the requirements for the award of Bachelor of '
    'Technology in Computer Science and Engineering, is a record of bonafide work done under my '
    'guidance. The contents of this Project work, in full or in parts, have neither been taken from '
    'any other source nor have been submitted to any other Institute or University for award of any '
    'degree or diploma and the same is certified.',
    space_after=24, first_line_indent=1.27
)

add_empty_lines(3)
add_normal_para('Dr. [Guide Name]')
add_normal_para('Internal Guide', space_after=24)

add_empty_lines(2)
add_normal_para('The thesis is satisfactory / unsatisfactory', space_after=24)

add_empty_lines(3)
p = doc.add_paragraph()
run1 = p.add_run('Internal Examiner')
set_font(run1, size=12)
run2 = p.add_run('                                                    External Examiner')
set_font(run2, size=12)

add_empty_lines(3)
add_centered_para('Approved by', space_after=24)
add_empty_lines(2)
add_centered_para('DEAN', size=12, bold=True)
add_centered_para('School of Computer Science and Engineering', size=12)

# ============================================================
# ABSTRACT
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('ABSTRACT', size=14, bold=True, underline=True, space_after=18)
add_empty_lines(1)

add_normal_para(
    'When a natural disaster strikes, the first few hours determine how many lives can be saved. '
    'Emergency responders currently face a fragmented information landscape where seismograph readings '
    'arrive without any visual confirmation of structural damage, satellite images show devastation '
    'without on-the-ground context, and social media posts describe chaos without precise physical '
    'measurements to back them up. Each of these data sources tells part of the story, but no single '
    'source provides a complete picture of what is happening on the ground.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'This work introduces a tri-modal fusion architecture that brings together three independent streams '
    'of disaster information\u2014environmental sensor metadata covering weather, seismic, storm, and '
    'hydrological conditions; post-disaster satellite imagery for structural damage assessment; and social '
    'media image-text pairs capturing the human dimension of the crisis\u2014through a pairwise cross-attention '
    'mechanism with adaptive gating. Rather than simply concatenating features from different sources, '
    'the system learns which signals from each modality to trust and how they should inform one another.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The architecture is organized around three specialized sub-models. The first, an AdaptiveIoTClassifier, '
    'processes a unified 32-dimensional sensor vector organized into four groups\u2014weather, storm, seismic, '
    'and hydrological\u2014each with its own learned confidence estimator that automatically suppresses '
    'irrelevant sensor groups. This sub-model achieves 97.64% classification accuracy across 63,527 samples '
    'drawn from six publicly available datasets. The second sub-model, an AdaptiveFusionClassifier, combines '
    'a BLIP Vision Transformer for image understanding with an XLM-RoBERTa encoder for multilingual text '
    'processing, connected through bidirectional cross-modal attention. It reaches 86.39% accuracy on the '
    'CrisisMMD humanitarian benchmark, improving upon support vector machine baselines by over twelve '
    'percentage points. The third sub-model adapts DeepLabV3+ with a ResNet101 backbone for satellite imagery '
    'from the xBD dataset, performing both pixel-level damage segmentation and compact 640-dimensional '
    'embedding extraction, trained with class-weighted cross-entropy loss on GPU-accelerated hardware.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'A TriFusionLayer integrates embeddings from all three sub-models through three pairwise cross-attention '
    'blocks and a learned modality gate, achieving 99.41% priority classification accuracy and a severity '
    'mean absolute error of 0.0398\u2014representing a 65.8% improvement over the crisis-only baseline. '
    'Ablation experiments across four modality configurations show that each added data source monotonically '
    'improves every metric, and a 30% modality dropout training strategy enables the system to function '
    'even when some data streams are unavailable. Noise sensitivity testing with Gaussian injection at '
    '5-20% levels confirms that classification performance degrades gracefully rather than collapsing. '
    'An explainability framework adapting gradient-weighted attention rollout for tri-modal Vision '
    'Transformer interpretability, backed by a four-tier fallback chain, ensures that every prediction '
    'comes with a human-readable justification.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Keywords: disaster management, multimodal fusion, cross-attention, environmental sensors, '
    'satellite imagery, social media analysis, deep learning, explainable AI',
    italic=True, space_after=12
)

# page number
add_centered_para('i', size=12, space_before=24)

# ============================================================
# ACKNOWLEDGEMENT
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('ACKNOWLEDGEMENT', size=14, bold=True, underline=True, space_after=18)
add_empty_lines(1)

add_normal_para(
    'I would like to express my sincere gratitude to Dr. [Guide Name], [Designation], School of '
    'Computer Science and Engineering, VIT-AP University, for the consistent guidance and encouragement '
    'throughout this project. The discussions on multimodal learning architectures and the practical '
    'challenges of deploying AI systems for disaster response shaped the direction of this work in '
    'ways that would not have been possible without that mentorship.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'I am grateful to the Chancellor, Vice Presidents, Vice Chancellor, and Dr. [Dean Name], Dean, '
    'School of Computer Science and Engineering, VIT-AP University, for fostering an academic '
    'environment that supports ambitious research endeavors and provides the computational resources '
    'necessary for deep learning experimentation.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'I extend my thanks to [Program Chair Name], Program Chair, and to all the faculty members of the '
    'Department of Computer Science and Engineering whose coursework in machine learning, computer '
    'vision, and natural language processing laid the foundation for this interdisciplinary project. '
    'The Kaggle GPU instances that enabled training of the satellite and crisis models were essential '
    'to achieving the results reported here.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'I owe a debt of gratitude to my parents for their unwavering support and patience during the '
    'long hours spent training models and debugging attention mechanisms. Their belief in the value '
    'of this work kept me motivated through the setbacks that are part of any research project.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Finally, I would like to thank my friends and peers who offered feedback on early versions of '
    'the fusion architecture, helped test the web dashboard, and contributed to productive discussions '
    'about the practical implications of automated disaster assessment systems.',
    space_after=24, first_line_indent=1.27
)

add_normal_para('Place: Amaravati', bold=True)
p = doc.add_paragraph()
run1 = p.add_run('Date:')
set_font(run1, size=12, bold=True)
run2 = p.add_run('\t\t\t\t\t\tMARURI SAI RAMA LINGA REDDY')
set_font(run2, size=12, bold=True)

add_centered_para('ii', size=12, space_before=24)

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('CONTENTS', size=14, bold=True, underline=True, space_after=18)
add_empty_lines(1)

toc_entries = [
    ('', 'Abstract', 'i'),
    ('', 'Acknowledgement', 'ii'),
    ('', 'Table of Contents', 'iii'),
    ('', 'List of Tables', 'iv'),
    ('', 'List of Figures', 'v'),
    ('', 'List of Abbreviations', 'vi'),
    ('1', 'Introduction', '1'),
    ('1.1', '    Background and Motivation', '1'),
    ('1.2', '    Problem Statement', '2'),
    ('1.3', '    Objectives', '3'),
    ('1.4', '    Scope of the Project', '3'),
    ('1.5', '    Organization of the Report', '4'),
    ('2', 'Literature Review', '5'),
    ('2.1', '    Single-Modality Disaster Analysis', '5'),
    ('2.2', '    Multimodal Fusion Approaches', '7'),
    ('2.3', '    Cross-Modal Attention Mechanisms', '8'),
    ('2.4', '    Explainability for Vision Transformers', '9'),
    ('2.5', '    Research Gap Analysis', '10'),
    ('3', 'Proposed Methodology', '11'),
    ('3.1', '    System Overview', '11'),
    ('3.2', '    AdaptiveIoTClassifier', '12'),
    ('3.3', '    AdaptiveFusionClassifier', '15'),
    ('3.4', '    DeepLabV3+ with Dual Feature Extraction', '18'),
    ('3.5', '    TriFusionLayer', '20'),
    ('3.6', '    Explainability Framework', '23'),
    ('3.7', '    Training Strategy', '25'),
    ('4', 'Experimental Setup', '27'),
    ('4.1', '    Datasets', '27'),
    ('4.2', '    Feature-to-Dimension Mapping', '29'),
    ('4.3', '    Implementation Details', '30'),
    ('4.4', '    Evaluation Metrics', '31'),
    ('5', 'Results and Discussion', '32'),
    ('5.1', '    IoT Model Performance', '32'),
    ('5.2', '    Crisis Model Performance', '35'),
    ('5.3', '    xBD Satellite Model Performance', '38'),
    ('5.4', '    Tri-Fusion Ablation Study', '40'),
    ('5.5', '    Noise Sensitivity Analysis', '42'),
    ('5.6', '    Discussion', '44'),
    ('6', 'Conclusion and Future Work', '47'),
    ('6.1', '    Summary of Contributions', '47'),
    ('6.2', '    Limitations', '48'),
    ('6.3', '    Future Work', '49'),
    ('', 'References', '51'),
    ('', 'Appendices', '55'),
]

for num, title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if num:
        run = p.add_run(f'{num}\t{title}')
    else:
        run = p.add_run(f'{title}')
    set_font(run, size=12, bold=(len(num) <= 1))
    # add dots and page number
    run2 = p.add_run(f'\t{page}')
    set_font(run2, size=12)

add_centered_para('iii', size=12, space_before=24)

# ============================================================
# LIST OF TABLES
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('LIST OF TABLES', size=14, bold=True, underline=True, space_after=18)

tables_list = [
    ('3.1', 'Learned Sensor Group Weights by Disaster Type', '14'),
    ('4.1', 'Mapping of Kaggle Dataset Features to the 32-Dimensional Unified Sensor Vector', '29'),
    ('4.2', 'Training Configuration for Each Sub-Model', '30'),
    ('5.1', 'IoT Model Per-Class Classification Results', '33'),
    ('5.2', 'Crisis Model Per-Class Test Results', '36'),
    ('5.3', 'Comparison with Baseline and State-of-the-Art Models on CrisisMMD', '37'),
    ('5.4', 'xBD Satellite Model Per-Class Results', '39'),
    ('5.5', 'Ablation Study: Modality Contribution to Tri-Fusion Performance', '41'),
    ('5.6', 'Noise Sensitivity Analysis Results', '43'),
    ('5.7', 'Tri-Fusion Training Convergence', '44'),
]

for num, title, page in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'Table {num}\t{title}\t{page}')
    set_font(run, size=12)

add_centered_para('iv', size=12, space_before=24)

# ============================================================
# LIST OF FIGURES
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('LIST OF FIGURES', size=14, bold=True, underline=True, space_after=18)

figures_list = [
    ('3.1', 'Overview of the Proposed TriFusion Architecture', '11'),
    ('3.2', 'AdaptiveIoTClassifier Architecture', '13'),
    ('3.3', 'AdaptiveFusionClassifier Architecture', '16'),
    ('3.4', 'DeepLabV3+ Modified Architecture for xBD', '19'),
    ('3.5', 'TriFusionLayer with Pairwise Cross-Attention', '21'),
    ('3.6', 'Gradient-Weighted Attention Rollout Pipeline', '24'),
    ('5.1', 'Learned Sensor Group Weights per Disaster Class', '33'),
    ('5.2', 'IoT Model Confusion Matrix', '34'),
    ('5.3', 'Per-Class ROC Curves for AdaptiveIoTClassifier', '34'),
    ('5.4', 'IoT Severity Regression Performance', '35'),
    ('5.5', 'Crisis Model Training History', '36'),
    ('5.6', 'Crisis Model Confusion Matrix', '37'),
    ('5.7', 'Crisis Model ROC Curves', '38'),
    ('5.8', 'xBD Training History on GPU', '39'),
]

for num, title, page in figures_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'Figure {num}\t{title}\t{page}')
    set_font(run, size=12)

add_centered_para('v', size=12, space_before=24)

# ============================================================
# LIST OF ABBREVIATIONS
# ============================================================
add_page_break()
add_empty_lines(1)
add_centered_para('LIST OF ABBREVIATIONS', size=14, bold=True, underline=True, space_after=18)

abbrevs = [
    ('AI', 'Artificial Intelligence'),
    ('AUC', 'Area Under the Curve'),
    ('BLIP', 'Bootstrapping Language-Image Pre-training'),
    ('CE', 'Cross-Entropy'),
    ('CLS', 'Classification Token'),
    ('CNN', 'Convolutional Neural Network'),
    ('DEM', 'Digital Elevation Model'),
    ('GIS', 'Geographic Information System'),
    ('GPU', 'Graphics Processing Unit'),
    ('IoT', 'Internet of Things'),
    ('MAE', 'Mean Absolute Error'),
    ('MHA', 'Multi-Head Attention'),
    ('MLP', 'Multi-Layer Perceptron'),
    ('mIoU', 'Mean Intersection over Union'),
    ('MSE', 'Mean Squared Error'),
    ('NDVI', 'Normalized Difference Vegetation Index'),
    ('NDWI', 'Normalized Difference Water Index'),
    ('NOAA', 'National Oceanic and Atmospheric Administration'),
    ('RGB', 'Red Green Blue'),
    ('RMS', 'Root Mean Square'),
    ('ROC', 'Receiver Operating Characteristic'),
    ('USGS', 'United States Geological Survey'),
    ('ViT', 'Vision Transformer'),
    ('XAI', 'Explainable Artificial Intelligence'),
    ('xBD', 'xView2 Building Damage'),
    ('XLM-R', 'Cross-lingual Language Model - RoBERTa'),
]

for abbr, full in abbrevs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{abbr}\t\t{full}')
    set_font(run, size=12)

add_centered_para('vi', size=12, space_before=24)

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_page_break()
add_heading_custom('Chapter 1', level=0)
add_heading_custom('Introduction', level=0)
add_empty_lines(1)

add_heading_custom('1.1 Background and Motivation', level=1)

add_normal_para(
    'Natural disasters remain among the most destructive forces affecting human societies. The year 2023 '
    'alone saw global economic losses exceeding 380 billion US dollars, with delayed situational '
    'awareness repeatedly identified as a contributing factor in preventable casualties. When a cyclone '
    'makes landfall or an earthquake shakes a populated region, the hours immediately following the event '
    '\u2014 often called the golden window \u2014 are the most critical for mounting an effective response. '
    'During this narrow timeframe, first responders must determine what has happened, where the worst '
    'damage lies, who needs help most urgently, and what resources should be deployed. These decisions '
    'are only as good as the information that feeds them.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The challenge is that disaster information arrives through multiple channels, each offering a '
    'different piece of the puzzle. Weather stations and seismographs provide precise physical measurements '
    '\u2014 wind speed, barometric pressure, ground acceleration, rainfall intensity \u2014 but these readings '
    'alone say nothing about whether buildings have collapsed or whether people are trapped. Satellite '
    'imagery can reveal the spatial extent of destruction across wide areas, showing which neighborhoods '
    'suffered the worst structural damage, but satellite passes may be delayed by 30 to 60 minutes after '
    'an event and cannot capture the human experience of the crisis. Social media posts from affected '
    'individuals provide raw, immediate accounts of what is happening on the ground \u2014 people reporting '
    'flooded streets, sharing photographs of damaged homes, calling for rescue \u2014 but these reports are '
    'noisy, unstructured, and lack the quantitative precision of sensor readings.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Existing disaster management systems overwhelmingly rely on a single one of these data sources. '
    'Earthquake early warning networks monitor seismographs but do not incorporate visual evidence of '
    'damage. Crisis informatics platforms mine Twitter and other social networks for on-the-ground reports '
    'but cannot cross-validate those reports against physical sensor readings. Satellite damage assessment '
    'tools produce building-level damage maps but operate in isolation from both sensor data and social '
    'media context. Each system captures complementary but incomplete information, and emergency managers '
    'are left to mentally integrate outputs from multiple disconnected tools under extreme time pressure.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'This project is motivated by the observation that if a neural architecture could learn to reason '
    'across all three data modalities simultaneously \u2014 asking, in effect, does the social media post '
    'about building collapse match what the satellite shows, and do the seismic readings support an '
    'earthquake strong enough to cause that level of damage? \u2014 the resulting situational assessment '
    'would be substantially more accurate and trustworthy than any single-source analysis.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('1.2 Problem Statement', level=1)

add_normal_para(
    'The fundamental problem addressed in this work is the lack of a unified computational framework '
    'that can ingest, align, and reason across heterogeneous disaster data streams in real time. '
    'Specifically, no prior system fuses all three modalities \u2014 environmental sensor metadata, '
    'satellite damage imagery, and social media image-text pairs \u2014 through a learned cross-attention '
    'mechanism that allows each data source to inform the interpretation of the others. Furthermore, '
    'existing multimodal systems typically require all inputs to be present or they fail entirely, which '
    'is impractical in real disaster scenarios where different data streams become available at different '
    'times after event onset.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('1.3 Objectives', level=1)

add_normal_para('The primary objectives of this project are:', space_after=6)

objectives = [
    'To design and implement a tri-modal fusion architecture that integrates environmental sensor '
    'metadata (weather, seismic, storm, hydrological), satellite damage imagery, and social media '
    'image-text pairs through pairwise cross-attention with adaptive gating.',

    'To develop an adaptive confidence-weighted sensor fusion mechanism where per-group learned '
    'confidence estimators automatically determine the relevance of each sensor group (weather, '
    'storm, seismic, hydro) for the detected disaster type.',

    'To implement graceful modality degradation so that the system produces useful output from '
    'whatever data streams are available, without requiring retraining or manual mode switching, '
    'using a 30% modality dropout training strategy.',

    'To adapt gradient-weighted attention rollout for tri-modal Vision Transformer explainability '
    'and implement a four-tier fallback chain that guarantees interpretable output regardless of '
    'model internals or computation failures.',

    'To conduct a comprehensive ablation study quantifying the contribution of each modality to '
    'the fused output, along with a noise sensitivity analysis validating robustness under '
    'simulated real-world sensor degradation.',
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(f'{i}. {obj}')
    set_font(run, size=12)

add_heading_custom('1.4 Scope of the Project', level=1)

add_normal_para(
    'This project encompasses the end-to-end design, implementation, training, and evaluation of the '
    'tri-modal fusion architecture. The environmental sensor sub-model is trained on historical metadata '
    'compiled from six publicly available datasets covering California wildfire conditions, tropical '
    'cyclone tracks, NOAA Atlantic hurricane records, global earthquake catalogs, Iran earthquake data, '
    'and Sri Lanka flood risk assessments. The crisis sub-model is trained and evaluated on the CrisisMMD '
    'dataset containing 8,079 image-text pairs from real disaster events. The satellite sub-model is '
    'trained on the xBD dataset with post-disaster satellite images across 12 disaster events.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The scope includes a web-based demonstration dashboard built with FastAPI and vanilla JavaScript '
    'that allows users to submit sensor readings, upload disaster images, and enter associated text to '
    'receive real-time tri-modal assessments with visual and textual explanations. The scope does not '
    'include deployment on real-time operational sensor networks, though the architecture is designed '
    'to accommodate such deployment.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('1.5 Organization of the Report', level=1)

add_normal_para(
    'This report is organized into six chapters. Chapter 1 introduces the problem, motivation, and '
    'objectives. Chapter 2 reviews related work spanning single-modality disaster analysis, multimodal '
    'fusion approaches, cross-modal attention mechanisms, and Vision Transformer explainability. '
    'Chapter 3 describes the proposed methodology in detail, covering each sub-model architecture, '
    'the tri-fusion layer, and the explainability framework. Chapter 4 presents the experimental setup '
    'including datasets, feature engineering, and training configurations. Chapter 5 reports and discusses '
    'experimental results, the ablation study, and the noise sensitivity analysis. Chapter 6 concludes '
    'the work, acknowledges limitations, and outlines directions for future research.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# CHAPTER 2: LITERATURE REVIEW
# ============================================================
add_page_break()
add_heading_custom('Chapter 2', level=0)
add_heading_custom('Literature Review', level=0)
add_empty_lines(1)

add_heading_custom('2.1 Single-Modality Disaster Analysis', level=1)

add_heading_custom('2.1.1 Social Media Analysis', level=2)

add_normal_para(
    'The CrisisMMD dataset, introduced by Alam, Ofli, and Imran in 2018, established one of the '
    'first large-scale benchmarks for multimodal crisis content classification. Their work showed '
    'that combining image and text features improves humanitarian categorization compared to '
    'text-only or image-only baselines, although their fusion strategy relied on simple feature '
    'concatenation rather than learned attention. CrisisNLP and CrisisLex focused specifically on '
    'text classification for crisis-related tweets, building lexicons and classification pipelines '
    'for filtering relevant messages from the noise of social media during emergencies.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Abavisani and colleagues in 2020 explored attention-based fusion for crisis tweets, demonstrating '
    'that attention mechanisms outperform concatenation for this task. However, their work remained '
    'within the social media modality and did not incorporate sensor data or satellite imagery. More '
    'recently, Mouzannar, Rizk, and Awad introduced Guided Cross-Attention with LLaVA-generated '
    'captions, pushing CrisisMMD accuracy to 93.92 percent. This represents the current state of the '
    'art on the CrisisMMD humanitarian task but focuses on maximizing single-task classification '
    'accuracy rather than producing embeddings suitable for downstream multi-modal fusion.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.1.2 Satellite Damage Assessment', level=2)

add_normal_para(
    'The xBD dataset, released as part of the xView2 challenge, provides the largest publicly '
    'available building damage assessment benchmark with pre- and post-disaster satellite imagery. '
    'Winning solutions in the challenge employed U-Net variants and segmentation architectures '
    'optimized for pixel-level damage classification across four damage levels. Weber and Kan applied '
    'change detection networks that compare pre- and post-disaster images to identify damaged structures. '
    'More recent approaches using optical-only configurations have pushed Mean IoU scores to between '
    '66.19 and 69.76 percent on the xBD benchmark through combinations of Dice and Focal loss '
    'functions, multi-scale feature aggregation, and aggressive augmentation pipelines. All of these '
    'approaches produce pixel-level segmentation outputs but lack any mechanism for integrating their '
    'findings with real-time sensor readings or social media reports.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.1.3 Environmental Sensor Monitoring', level=2)

add_normal_para(
    'Earthquake early warning systems such as the ElarmS system developed by Allen demonstrate the '
    'value of sensor-based approaches for rapid disaster detection. Flood monitoring networks like '
    'those described by Basha and Rus use distributed sensor deployments to detect rising water levels '
    'and trigger alerts. The Sakaki earthquake detection system on Twitter showed that social sensors '
    'could complement physical sensors by detecting events through the volume and content of user posts. '
    'However, all of these systems operate independently. A seismic alert system does not check whether '
    'satellite imagery confirms building damage, and a flood sensor network does not cross-reference '
    'social media reports to determine whether affected populations have evacuated.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.2 Multimodal Fusion Approaches', level=1)

add_normal_para(
    'Multimodal fusion in disaster contexts remains surprisingly limited in the published literature. '
    'Ofli, Alam, and Imran combined satellite and social media data for damage assessment but used '
    'feature concatenation rather than learned attention, treating both modalities equally regardless '
    'of their information content for a given sample. Rudner and colleagues proposed Multi3Net for '
    'satellite multimodal fusion combining multiresolution, multisensor, and multitemporal imagery, '
    'but their work focused on a single disaster type (flooding) and did not incorporate social media '
    'or sensor data.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'To the best of our knowledge, no prior published work fuses all three modalities \u2014 '
    'environmental sensor metadata, satellite imagery, and social media image-text pairs \u2014 through '
    'a learned cross-attention mechanism with adaptive gating. This gap in the literature motivates '
    'the architecture presented in this thesis.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.3 Cross-Modal Attention Mechanisms', level=1)

add_normal_para(
    'The transformer architecture introduced by Vaswani and colleagues in 2017 revolutionized sequence '
    'modeling through its self-attention mechanism. Cross-attention, where queries from one modality '
    'attend to keys and values from another, has since become the dominant approach for multimodal '
    'fusion in vision-language tasks. ViLBERT proposed co-attentional transformer layers where visual '
    'and linguistic streams interact through cross-modal attention. LXMERT similarly used cross-modality '
    'encoder layers for learning representations that span images and text.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'BLIP, the Bootstrapping Language-Image Pre-training framework, demonstrated that a unified '
    'vision-language model pre-trained on large-scale image-caption pairs could achieve strong '
    'performance across understanding and generation tasks. Our work builds on BLIP by using its '
    'Vision Transformer as the visual backbone of the crisis sub-model, leveraging its pre-trained '
    'representations for disaster image understanding.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.4 Explainability for Vision Transformers', level=1)

add_normal_para(
    'Standard Grad-CAM, which works well for convolutional neural networks, fails on Vision '
    'Transformers because only the classification token is used for prediction, leaving gradients '
    'on individual patch tokens near zero. Chefer, Gur, and Wolf addressed this limitation in 2021 '
    'by proposing a method that propagates gradient-weighted self-attention maps across all transformer '
    'layers, producing relevancy scores that accurately reflect which input patches influenced the '
    'prediction. Their approach has become the standard tool for Vision Transformer interpretability. '
    'Our work adapts this technique to the specific demands of tri-modal disaster assessment, '
    'extending it with a multi-tier fallback chain and integrating the visual explanations with '
    'natural language briefings generated through GPT-4o.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('2.5 Research Gap Analysis', level=1)

add_normal_para(
    'The literature review reveals several gaps that this work aims to address. First, existing '
    'disaster analysis systems operate on single data modalities, missing the complementary information '
    'available from other sources. Second, the few multimodal approaches that exist rely on simple '
    'concatenation rather than learned cross-modal attention. Third, no prior architecture handles '
    'missing modalities gracefully, requiring all inputs or failing entirely. Fourth, explainability '
    'techniques developed for standard Vision Transformers have not been adapted for multi-modal '
    'disaster assessment pipelines. This project addresses all four gaps through its tri-modal fusion '
    'architecture with adaptive gating, modality dropout training, and four-tier explainability chain.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# CHAPTER 3: PROPOSED METHODOLOGY
# ============================================================
add_page_break()
add_heading_custom('Chapter 3', level=0)
add_heading_custom('Proposed Methodology', level=0)
add_empty_lines(1)

add_heading_custom('3.1 System Overview', level=1)

add_normal_para(
    'The proposed architecture processes three independent data streams through specialized sub-models, '
    'each producing fixed-dimensional embeddings that are fused through pairwise cross-attention with '
    'adaptive gating. The first stream consists of environmental and seismological sensor metadata '
    'encoded as a 32-dimensional feature vector. The second stream consists of social media image-text '
    'pairs from disaster-affected areas. The third stream consists of post-disaster satellite imagery '
    'at 512 by 512 pixel resolution.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Each data stream is processed by a dedicated sub-model: the AdaptiveIoTClassifier produces a '
    '128-dimensional embedding from sensor data, the AdaptiveFusionClassifier produces a '
    '1024-dimensional embedding from image-text pairs, and a modified DeepLabV3+ produces a '
    '640-dimensional embedding from satellite imagery. These three embeddings are then fed into '
    'a TriFusionLayer that performs three pairwise cross-attention operations (crisis-IoT, '
    'crisis-satellite, IoT-satellite) followed by adaptive gating, producing five output predictions: '
    'severity score, priority level, disaster type, population impact, and resource allocation needs.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'A note on terminology is warranted here. Throughout this report, the term IoT is used as a '
    'convenient shorthand for the environmental and seismological metadata sub-model. The training data '
    'for this sub-model consists of historical weather records, seismological catalogs, tropical storm '
    'track databases, and a synthetic flood risk dataset, which are archival metadata rather than '
    'real-time sensor streams from MQTT or LoRaWAN devices. The architectural design, however, is '
    'sensor-agnostic: the 32-dimensional input vector and adaptive confidence weighting can be directly '
    'applied to live sensor deployments.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.2 AdaptiveIoTClassifier (Environmental Sensor Sub-Model)', level=1)

add_normal_para(
    'The environmental sensor sub-model processes a unified 32-dimensional feature vector encoding '
    'four sensor groups \u2014 weather, storm, seismic, and hydrological \u2014 each occupying 8 dimensions '
    'with consistent normalization to the range of zero to one.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.2.1 Sensor Group Encoding', level=2)

add_normal_para(
    'The 32-dimensional input is split into four groups of 8 features each. Each group is processed '
    'by an independent encoder consisting of two linear layers with ReLU activation, batch normalization, '
    'and dropout, mapping each 8-dimensional input to a 128-dimensional representation. The weather '
    'group (dimensions 0 through 7) encodes precipitation, temperature extremes, wind speed, drought '
    'index, and cyclic month encoding. The storm group (dimensions 8 through 15) captures wind intensity, '
    'pressure anomaly, geographic coordinates, storm category, and temporal features. The seismic group '
    '(dimensions 16 through 23) represents earthquake depth, RMS error, station count, azimuth gap, '
    'and magnitude. The hydrological group (dimensions 24 through 31) contains elevation, river '
    'proximity, rainfall measures, drainage, vegetation indices, and flood history.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.2.2 Confidence-Weighted Fusion', level=2)

add_normal_para(
    'A distinguishing feature of this sub-model is the per-group confidence estimation mechanism. '
    'Each sensor group has a dedicated SensorConfidenceEstimator network that takes the '
    '128-dimensional group encoding as input and produces a scalar confidence value between zero and '
    'one through a three-layer architecture (128 to 64 to 32 to 1) with a sigmoid output. The four '
    'confidence values are then normalized to produce weights that sum to one. Each group encoding is '
    'multiplied by its corresponding weight before being passed to the cross-group attention stage.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'This mechanism serves a critical function: when a sensor group carries no useful information '
    'for the current disaster type \u2014 for instance, seismic features during a flood event \u2014 the '
    'confidence estimator learns to assign near-zero weight to that group, preventing noise from '
    'irrelevant sensors from corrupting the final prediction. The learned weights demonstrate '
    'physically meaningful behavior: storm detection assigns 99.9 percent weight to the storm sensor '
    'group, earthquake detection assigns 82.8 percent to seismic sensors, flood detection assigns '
    '76.7 percent to hydrological sensors, and fire detection assigns 75.8 percent to weather sensors.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.2.3 Cross-Group Multi-Head Attention', level=2)

add_normal_para(
    'After confidence weighting, the four group feature vectors are stacked and fed into a multi-head '
    'self-attention layer with 4 heads and an embedding dimension of 128. This cross-group attention '
    'allows sensor groups to exchange information, which is particularly valuable for cascading '
    'disasters where one hazard triggers another. For example, seismic activity can reinforce '
    'hydrological signals to detect earthquake-triggered tsunamis, or high winds combined with low '
    'pressure can distinguish a hurricane from an isolated thunderstorm. The attended features are '
    'mean-pooled across the four groups to produce a single 128-dimensional global embedding.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.2.4 Multi-Task Output Heads', level=2)

add_normal_para(
    'Four parallel output heads operate on the shared 128-dimensional embedding. The disaster type '
    'head performs 5-class classification over fire, storm, earthquake, flood, and unknown categories. '
    'The severity head produces a scalar regression value between zero and one indicating overall '
    'disaster intensity. The risk detail head outputs a 4-dimensional vector of per-hazard risk scores. '
    'The casualty risk head produces a scalar estimate of human casualty likelihood. The combined '
    'training loss is the sum of cross-entropy loss for the type classification and mean squared error '
    'losses for the three regression tasks.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.3 AdaptiveFusionClassifier (Crisis Sub-Model)', level=1)

add_normal_para(
    'The crisis sub-model is responsible for understanding social media content from disaster-affected '
    'areas. It processes image-text pairs and produces both humanitarian category predictions and '
    'compact 1024-dimensional embeddings for downstream tri-modal fusion.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.3.1 Feature Extraction', level=2)

add_normal_para(
    'Visual features are extracted using a BLIP Vision Transformer (ViT-B/16) pre-trained on COCO '
    'captions. Input images are resized to 224 by 224 pixels and normalized using ImageNet statistics. '
    'The ViT processes each image into 196 patch tokens plus one classification token, all with '
    '768-dimensional representations. The CLS token embedding serves as the global image representation. '
    'Text features are extracted using XLM-RoBERTa, a cross-lingual transformer encoder supporting '
    '100 languages, which is important because disaster-related tweets are often multilingual. '
    'Input text is tokenized with a maximum length of 128 tokens, and the CLS token from the final '
    'layer provides a 768-dimensional sentence embedding.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.3.2 Adaptive Modality Weighting', level=2)

add_normal_para(
    'Rather than treating image and text features as equally important for every input sample, the '
    'model includes separate ConfidenceEstimator networks for each modality. Each estimator takes '
    'the 768-dimensional raw feature as input and produces a scalar confidence through a three-layer '
    'network (768 to 256 to 128 to 1 with sigmoid). The two confidence scores are normalized to sum '
    'to one, creating input-dependent modality weights. When an image is blurry or uninformative, '
    'the visual confidence decreases and the model relies more heavily on the text. When a tweet is '
    'generic, the model shifts weight toward the visual content.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.3.3 Bidirectional Cross-Modal Attention', level=2)

add_normal_para(
    'Both modalities are projected to 512 dimensions through dedicated projection layers with layer '
    'normalization, ReLU activation, and dropout. The projected and weighted representations then '
    'participate in bidirectional cross-modal attention with 8 heads and a dimension of 512. In the '
    'first direction, visual features serve as queries while text features provide keys and values, '
    'allowing the image representation to attend to relevant textual context. In the second direction, '
    'text features query the visual representation. The two cross-attended outputs are concatenated '
    'to form the final 1024-dimensional crisis embedding that is passed downstream to the '
    'TriFusionLayer.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.4 DeepLabV3+ with Dual Feature Extraction (xBD Sub-Model)', level=1)

add_normal_para(
    'The satellite sub-model adapts the DeepLabV3+ segmentation architecture with a ResNet101 encoder '
    'backbone to process 512 by 512 pixel post-disaster satellite images from the xBD dataset. This '
    'sub-model has a dual purpose: producing pixel-level damage segmentation maps and extracting '
    'compact embeddings that capture global damage characteristics for fusion.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.4.1 Loss Function and Training', level=2)

add_normal_para(
    'The xBD dataset exhibits severe class imbalance, with no-damage pixels comprising over 85 percent '
    'of annotated regions. The model is trained with class-weighted cross-entropy loss. In the final '
    'GPU-accelerated configuration, uniform class weights performed well when combined with an extensive '
    'augmentation pipeline including horizontal and vertical flips, random 90-degree rotations, '
    'shift-scale-rotate transformations, brightness-contrast adjustments, and Gaussian noise injection. '
    'A cosine annealing learning rate schedule and gradient clipping at 1.0 ensure stable convergence. '
    'A differential learning rate strategy is employed: the pre-trained ResNet101 encoder trains at '
    '3 times 10 to the negative 5, the decoder at 1 times 10 to the negative 4, and projection heads '
    'at 2 times 10 to the negative 4.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.4.2 Embedding Extraction', level=2)

add_normal_para(
    'Two embedding pathways operate in parallel. The satellite embedding pathway applies adaptive '
    'average pooling to the encoder final feature map (2048 by 16 by 16) followed by a linear '
    'projection to produce a 512-dimensional vector capturing global damage characteristics. The '
    'regional statistics pathway processes concatenated decoder features and damage probability maps '
    'through a RegionalStatsModule with convolutional fusion, 4 by 4 pooling, and fully connected '
    'layers to produce a 128-dimensional spatial statistics embedding. These two vectors are '
    'concatenated to form the final 640-dimensional satellite embedding.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.5 TriFusionLayer', level=1)

add_normal_para(
    'The TriFusionLayer is the central component that integrates embeddings from all three sub-models '
    'through pairwise cross-attention and adaptive gating.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.5.1 Modality Projection', level=2)

add_normal_para(
    'Each sub-model embedding is first projected to a common 256-dimensional space through dedicated '
    'multi-layer perceptrons. The crisis embedding (1024 dimensions) passes through a 1024 to 512 to '
    '256 projection. The IoT embedding (128 dimensions) passes through a 128 to 512 to 256 projection. '
    'The satellite embedding (640 dimensions) passes through a 640 to 512 to 256 projection. These '
    'projections use layer normalization, GELU activation, and dropout.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.5.2 Pairwise Cross-Attention', level=2)

add_normal_para(
    'Three bidirectional cross-attention operations are performed, each with 4 attention heads. '
    'The first operation allows crisis and IoT embeddings to attend to each other, producing two '
    '256-dimensional attended representations. The second allows crisis and satellite embeddings to '
    'interact. The third allows IoT and satellite embeddings to exchange information. This pairwise '
    'design produces six 256-dimensional attended representations in total, capturing all possible '
    'inter-modality interactions.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.5.3 Adaptive Modality Gating', level=2)

add_normal_para(
    'A learned gating network takes the concatenation of the three projected embeddings as input and '
    'produces softmax-normalized weights over the three modalities. Critically, when a modality is '
    'missing at inference time, its gate weight is set to negative infinity before the softmax '
    'operation, ensuring that the remaining modalities share the full weight. This mechanism is what '
    'enables graceful degradation: the model does not need to know in advance which modalities will '
    'be available.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.5.4 Output Heads', level=2)

add_normal_para(
    'The gate-weighted combination of projected and cross-attended features (1536 dimensions total) '
    'passes through a shared multi-layer perceptron (1536 to 512 to 256) followed by five parallel '
    'output heads. The severity head produces a scalar in the zero-to-one range through sigmoid '
    'activation. The priority head outputs 4-class logits (low, medium, high, critical). The '
    'disaster type head outputs 5-class logits. The population impact head estimates affected '
    'population fraction through a two-layer MLP with sigmoid. The resource needs head outputs a '
    '4-dimensional vector (water, medical, rescue, shelter) each between zero and one.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.6 Explainability Framework', level=1)

add_heading_custom('3.6.1 Gradient-Weighted Attention Rollout', level=2)

add_normal_para(
    'Standard Grad-CAM fails on Vision Transformers because only the CLS token is used for '
    'classification, resulting in near-zero gradients on individual patch tokens. To address this, '
    'the system adapts the gradient-weighted attention rollout technique introduced by Chefer, Gur, '
    'and Wolf. The procedure involves a forward pass through the BLIP ViT encoder with attention '
    'output enabled, capturing attention matrices for all 12 transformer layers. A backward pass '
    'from the predicted class logit captures gradient tensors through registered hooks. For the '
    'last 4 layers, gradient weighting is applied: attention matrices are element-wise multiplied '
    'with the positive-clamped gradients, retaining only features that increase the predicted class '
    'score. The weighted attention matrices are then rolled out across layers with a residual '
    'connection (50 percent attention plus 50 percent identity), producing a relevancy map from '
    'each patch to the CLS token.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The adaptation differs from the original method in two respects. First, gradient weighting is '
    'restricted to the last 4 layers rather than all 12, which empirically reduces noise in disaster '
    'imagery containing large uniform regions such as sky and water. Second, a 40th-percentile '
    'threshold with Gaussian smoothing is applied to produce field-deployable heatmaps rather than '
    'raw relevancy scores.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.6.2 Four-Tier XAI Fallback Chain', level=2)

add_normal_para(
    'To guarantee interpretable output regardless of computation failures, a cascading fallback '
    'is implemented. The first tier uses the full gradient-weighted attention rollout. If gradients '
    'fail, the second tier falls back to pure attention rollout without gradient weighting. If '
    'attention is unavailable, the third tier uses input gradient saliency. If all methods fail, '
    'the fourth tier returns a graceful empty result with a message indicating that explainability '
    'was not available. This chain ensures that the system always produces some level of explanation.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.7 Training Strategy', level=1)

add_heading_custom('3.7.1 Modality Dropout', level=2)

add_normal_para(
    'During tri-fusion training, IoT and satellite embeddings are independently dropped with 30 '
    'percent probability and replaced with learned default embeddings. This forces the model to '
    'develop robust internal representations that do not collapse when a modality is absent. At '
    'inference time, the same mechanism handles genuinely missing data by substituting the learned '
    'defaults, with the modality gate automatically redistributing weight to available modalities.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('3.7.2 Zero-Input Disaster Type Detection', level=2)

add_normal_para(
    'Both the IoT and crisis sub-models independently infer disaster type from raw data. The IoT '
    'model classifies the disaster type directly from the 32-dimensional sensor vector. The crisis '
    'model infers disaster type through a two-stage process: BLIP generates an image caption, and '
    'keyword scoring against disaster-specific vocabulary produces type probabilities. The '
    'TriFusionLayer reconciles both predictions. The system never requires manual hazard selection, '
    'saving critical minutes at incident onset.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# CHAPTER 4: EXPERIMENTAL SETUP
# ============================================================
add_page_break()
add_heading_custom('Chapter 4', level=0)
add_heading_custom('Experimental Setup', level=0)
add_empty_lines(1)

add_heading_custom('4.1 Datasets', level=1)

add_heading_custom('4.1.1 Environmental Sensor Metadata', level=2)

add_normal_para(
    'The environmental sensor sub-model was trained on 63,527 samples compiled from six publicly '
    'available Kaggle datasets. The California Weather and Fire Prediction dataset contributed '
    '14,988 rows of weather station readings linked to wildfire occurrence. The Tropical Cyclone '
    'Tracks dataset provided 59,228 records of historical cyclone positions and intensities. The '
    'NOAA Atlantic Hurricane dataset added 22,705 hurricane observations spanning 1975 to 2024. '
    'Two earthquake datasets \u2014 the CORGIS Earthquakes dataset from Virginia Tech and a dataset of '
    'earthquakes in Iran covering 1996 to 2022 \u2014 provided seismological records with magnitude, '
    'depth, and station metadata. The Sri Lanka Flood Risk and Inundation dataset contributed '
    '25,000 records of flood risk assessments.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'After preprocessing and subsampling, the final class distribution was: fire (4,971 samples), '
    'storm (13,162), earthquake (18,394), flood (25,000), and unknown (2,000). All features were '
    'encoded into a unified 32-dimensional vector with cyclic sine-cosine encoding for temporal '
    'features.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'An important data integrity measure was the exclusion of derived risk scores from the model '
    'input features. The Sri Lanka dataset contains target-correlated metadata such as '
    'flood_risk_score and infrastructure_score that would constitute data leakage if included as '
    'inputs. The model trains exclusively on primary physical variables, with derived scores used '
    'only for generating regression labels.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('4.1.2 Crisis Social Media', level=2)

add_normal_para(
    'The CrisisMMD dataset contains 8,079 image-text pairs from real disaster events, categorized '
    'into five humanitarian classes: affected individuals, infrastructure and utility damage, not '
    'humanitarian, other relevant information, and rescue or donation effort. The data was split '
    'into 6,126 training, 998 validation, and 955 test samples.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('4.1.3 Satellite Imagery', level=2)

add_normal_para(
    'The xBD dataset provides post-disaster satellite images at 512 by 512 pixel resolution across '
    'four damage classes (no-damage, minor, major, destroyed) spanning 12 disaster events including '
    'hurricanes, earthquakes, wildfires, flooding, tsunamis, and volcanic eruptions. After filtering '
    'and preprocessing, 2,684 samples were used with an 80/20 train-validation split.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('4.1.4 Tri-Fusion Training Data', level=2)

add_normal_para(
    'The fusion layer was trained on 13,608 samples from the CrisisMMD humanitarian split with '
    'synthetic IoT embeddings generated through the pre-trained AdaptiveIoTClassifier and real '
    'crisis embeddings, using an 80/20 train-validation split (10,886 training, 2,722 validation).',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('4.2 Implementation Details', level=1)

add_normal_para(
    'All models were implemented in PyTorch. The IoT, crisis, and xBD sub-models were trained on '
    'Kaggle GPU instances equipped with NVIDIA Tesla P100 GPUs with 16 gigabytes of memory. The '
    'tri-fusion layer was trained on CPU due to hardware scheduling constraints.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The IoT model trained for 40 epochs with a batch size of 256, using AdamW optimizer with '
    'learning rate 0.001 and weight decay 0.0001, and cosine annealing learning rate scheduling. '
    'The crisis model trained for 10 epochs with batch size 32 and a lower learning rate of '
    '0.00002 appropriate for fine-tuning pre-trained transformer encoders. The xBD model trained '
    'for 30 epochs with batch size 4 on GPU, achieving improved convergence in approximately 1.5 '
    'hours. The tri-fusion layer trained for 40 epochs with batch size 128.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('4.3 Evaluation Metrics', level=1)

add_normal_para(
    'Classification performance was evaluated using accuracy, precision, recall, F1-score (per-class '
    'and macro-averaged), and area under the ROC curve. Segmentation performance for the xBD model '
    'was measured using Intersection over Union (IoU) and F1-score per damage class, along with '
    'mean IoU across all classes. Regression tasks were evaluated using mean absolute error and, '
    'where applicable, R-squared coefficient of determination. The ablation study reported all five '
    'output metrics across four modality configurations to quantify each modality\'s contribution.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# CHAPTER 5: RESULTS AND DISCUSSION
# ============================================================
add_page_break()
add_heading_custom('Chapter 5', level=0)
add_heading_custom('Results and Discussion', level=0)
add_empty_lines(1)

add_heading_custom('5.1 IoT Model Performance', level=1)

add_normal_para(
    'The AdaptiveIoTClassifier achieves an overall accuracy of 97.64 percent with a macro F1-score '
    'of 89.99 percent across 63,527 samples. The model classifies storm, earthquake, and flood '
    'categories with perfect F1-scores of 1.000, while fire achieves 0.843 and the unknown category '
    'reaches 0.657. All classes achieve area under the ROC curve above 0.987, with a macro-average '
    'AUC of 0.996.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The perfect scores for storm, earthquake, and flood warrant careful interpretation. These results '
    'arise from three factors working together. Storm datasets contain high wind speed and low pressure '
    'values that are physically absent in other disaster categories. Earthquake datasets occupy a '
    'distinct subspace defined by magnitude, depth, and RMS values. The Sri Lanka flood data, being '
    'synthetically generated through rule-based simulations, exhibits clean separation from other '
    'categories. The fire-unknown confusion stems from overlapping weather-group features, as both '
    'categories share similar temperature and precipitation patterns.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Severity regression achieves a mean absolute error of 0.0452 and an R-squared value of 0.745, '
    'indicating that the model captures roughly three-quarters of the variance in disaster severity '
    'from sensor readings alone.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The learned sensor group weights provide strong evidence that the confidence estimation mechanism '
    'captures physically meaningful relationships. Storm detection relies 99.9 percent on the storm '
    'sensor group, earthquake detection relies 82.8 percent on seismic sensors, flood detection '
    'relies 76.7 percent on hydrological sensors, and fire detection relies 75.8 percent on weather '
    'sensors. This interpretability is a significant advantage for deployment in settings where '
    'operators need to understand and trust the system.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.2 Crisis Model Performance', level=1)

add_normal_para(
    'The AdaptiveFusionClassifier achieves a test accuracy of 86.39 percent and a validation F1-score '
    'of 87.48 percent on the CrisisMMD benchmark, representing an improvement of over 12 percentage '
    'points compared to logistic regression and SVM baselines. The strongest per-class performance '
    'comes from the not-humanitarian category (F1 = 0.89) and infrastructure damage (F1 = 0.88), '
    'while the affected-individuals class achieves the lowest F1-score of 0.43, attributable to its '
    'very low support of only 9 test samples.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The adaptive modality confidence scores remain balanced throughout training, with vision '
    'confidence settling around 0.50 and text confidence around 0.45. This balance confirms that '
    'the learned weighting mechanism avoids modality collapse, where one modality would dominate '
    'and the other would be ignored.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Compared to the concurrent work by Mouzannar and colleagues, who achieve 93.92 to 94.02 '
    'percent using Guided Cross Attention with LLaVA-generated captions, our crisis model is '
    'approximately 7.5 percentage points lower. However, the comparison is not entirely direct '
    'because our crisis model is designed to produce compact 1024-dimensional embeddings optimized '
    'for downstream tri-modal fusion, not to maximize standalone classification accuracy.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.3 xBD Satellite Model Performance', level=1)

add_normal_para(
    'The modified DeepLabV3+ achieves a best validation Mean IoU of 0.4276 at epoch 27 and a Mean '
    'F1-score of 0.5221. The no-damage class achieves near-perfect segmentation with an IoU of '
    '0.9774, while minority damage classes remain challenging: minor damage at 0.1595 IoU, major '
    'damage at 0.2647, and destroyed at 0.3087.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'This result represents a 52.6 percent improvement in Mean IoU compared to the initial CPU-trained '
    'configuration (mIoU = 0.2802), achieved through GPU-accelerated training with cosine annealing '
    'and an extensive augmentation pipeline. The GPU training completed 30 epochs in approximately '
    '1.5 hours at 3 minutes per epoch. The remaining gap with state-of-the-art methods (66 to 70 '
    'percent mIoU) is attributable to the extreme class imbalance and single-scale inference '
    'at 512 by 512 resolution.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Importantly, the satellite sub-model primary role in this architecture is embedding extraction '
    'for fusion, not standalone segmentation. The ablation study confirms that satellite embeddings '
    'contribute a 29.8 percentage point improvement in priority accuracy and a 64.4 percent '
    'reduction in severity MAE when added to the crisis-only baseline, demonstrating that the '
    'embeddings capture meaningful damage information even though pixel-level segmentation on '
    'minority classes has room for improvement.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.4 Tri-Fusion Ablation Study', level=1)

add_normal_para(
    'The ablation study constitutes the central experimental result of this work, evaluating '
    'performance across four modality configurations on 2,722 validation samples.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The crisis-only baseline achieves a severity MAE of 0.1165 and priority accuracy of 68.96 '
    'percent. Adding IoT sensor data reduces severity MAE by 15.1 percent to 0.0989 and improves '
    'priority accuracy by 3.41 percentage points to 72.37 percent. Adding satellite imagery alone '
    'has a much larger impact, reducing severity MAE by 64.4 percent to 0.0415 and boosting priority '
    'accuracy by 29.8 percentage points to 98.75 percent. The full tri-modal configuration achieves '
    'the best results across all metrics: severity MAE of 0.0398 (a 65.8 percent reduction from '
    'baseline), priority accuracy of 99.41 percent, population impact MAE of 0.0248 (71.7 percent '
    'reduction), and resource MAE of 0.0161 (65.6 percent reduction). Disaster type classification '
    'reaches 100 percent accuracy across all configurations.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'These results establish three important findings. First, satellite imagery is the strongest '
    'single contributor because it provides direct physical evidence of structural damage. Second, '
    'IoT sensor data provides incremental but meaningful improvement through environmental context. '
    'Third, each modality\'s addition monotonically improves every metric, confirming genuine '
    'complementarity rather than redundancy.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.5 Noise Sensitivity Analysis', level=1)

add_normal_para(
    'To validate that the perfect F1-scores for storm, earthquake, and flood reflect genuine '
    'discriminative capacity rather than overfitting to clean data, Gaussian noise was injected '
    'into the 32-dimensional sensor vectors at inference time with standard deviations of 0.05, '
    '0.10, and 0.20 (corresponding to 5, 10, and 20 percent of the feature range).',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'At 5 percent noise, the model retains over 99 percent of its clean performance. At 10 percent '
    'noise, storm, earthquake, and flood F1-scores remain above 0.98, confirming that inter-class '
    'separation reflects genuinely distinct physical signatures rather than artifacts of synthetic '
    'regularity. At 20 percent noise, representing severe sensor degradation well beyond typical '
    'operational conditions, the model still achieves 87.3 percent overall accuracy. The fire class '
    'shows the largest degradation due to its inherent feature overlap with the unknown category.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The confidence-weighted fusion mechanism contributes to this noise robustness. When noise '
    'corrupts one sensor group, the learned confidence estimators naturally down-weight the '
    'corrupted group contribution, providing an implicit noise-filtering effect that supplements '
    'the discriminative power of the feature representations themselves.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.6 Discussion', level=1)

add_heading_custom('5.6.1 Modality Complementarity', level=2)

add_normal_para(
    'The ablation study reveals a clear hierarchy in how different modalities contribute to the '
    'fused assessment. Satellite imagery contributes the most to priority and severity estimation '
    'because it provides direct physical evidence of structural damage that correlates strongly '
    'with response urgency. Environmental sensor metadata contributes contextual information about '
    'hazard conditions that is less directly observable from imagery but enriches the model '
    'understanding of cascading hazards. Social media provides the human dimension \u2014 who is '
    'affected, what help is being requested, how severe the perceived impact is \u2014 that neither '
    'sensors nor satellites can capture.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('5.6.2 Contextualizing the Priority Accuracy', level=2)

add_normal_para(
    'The 99.41 percent priority accuracy contrasts with real-world automated dispatch systems that '
    'typically report 85 to 87 percent accuracy on unstructured metadata. This gap reflects the '
    'fundamental advantage of tri-modal redundancy. In single-modality systems, an ambiguous social '
    'media post provides weak evidence for priority assignment. In the tri-modal system, the same '
    'post is cross-validated against simultaneous sensor readings and satellite visual evidence of '
    'damage. This multi-source confirmation substantially reduces classification ambiguity. The '
    'near-perfect accuracy is a direct consequence of the pairwise cross-attention design that '
    'enables each modality to resolve ambiguities in the others, as confirmed by the monotonic '
    'ablation results.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# CHAPTER 6: CONCLUSION AND FUTURE WORK
# ============================================================
add_page_break()
add_heading_custom('Chapter 6', level=0)
add_heading_custom('Conclusion and Future Work', level=0)
add_empty_lines(1)

add_heading_custom('6.1 Summary of Contributions', level=1)

add_normal_para(
    'This project developed and evaluated a tri-modal fusion architecture for real-time disaster '
    'intelligence that brings together environmental sensor metadata, satellite damage imagery, and '
    'social media image-text pairs through pairwise cross-attention with adaptive gating. The system '
    'achieves 99.41 percent priority classification accuracy and a severity mean absolute error of '
    '0.0398, representing a 65.8 percent improvement over the crisis-only baseline.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The key technical contributions include the adaptive confidence-weighted sensor fusion mechanism '
    'that automatically determines sensor group relevance for each disaster type, the graceful '
    'modality degradation enabled by 30 percent dropout training, the adaptation of gradient-weighted '
    'attention rollout for tri-modal Vision Transformer explainability, and the comprehensive '
    'evaluation through ablation study and noise sensitivity analysis that validates both the '
    'complementarity of the three modalities and the robustness of the approach under simulated '
    'real-world conditions.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The ablation study demonstrates that each added modality monotonically improves every evaluation '
    'metric, with satellite imagery providing the largest single contribution (29.8 percentage point '
    'priority improvement) and IoT sensor data adding incremental but meaningful gains (3.4 percentage '
    'points). The noise sensitivity analysis confirms that the underlying discriminative capacity '
    'extends well beyond the clean training distributions, with the model retaining over 98 percent '
    'of its F1-score for storm, earthquake, and flood categories even under 10 percent Gaussian '
    'noise injection.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('6.2 Limitations', level=1)

add_normal_para(
    'Several limitations of this work should be acknowledged. The environmental sensor sub-model is '
    'trained on historical archival metadata from public datasets rather than real-time IoT sensor '
    'streams. While the architecture is designed to be sensor-agnostic, real-time deployment introduces '
    'challenges including packet loss, battery-induced sensor drift, variable sampling rates, and '
    'multi-device calibration that are absent from the training data.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The training data includes a synthetic flood risk dataset that exhibits rule-based regularity, '
    'which partially explains the perfect classification scores for the flood category. The '
    'tri-fusion layer is trained with synthetic pairing of sub-model embeddings due to the absence '
    'of a naturally paired tri-modal disaster dataset. While the results are promising, performance '
    'on naturally co-occurring multi-modal data may differ.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'The xBD satellite model mIoU of 0.4276, while substantially improved over the initial '
    'configuration, remains below current state-of-the-art segmentation methods on the xBD benchmark. '
    'The tri-fusion layer was trained on CPU rather than GPU, which constrained batch size and '
    'training time exploration.',
    space_after=8, first_line_indent=1.27
)

add_heading_custom('6.3 Future Work', level=1)

add_normal_para(
    'Several directions for future work emerge from this project. First, the xBD satellite model can '
    'be improved through larger batch sizes, extended training beyond 100 epochs, multi-scale training '
    'and test-time augmentation, and copy-paste augmentation for minority damage classes. The target '
    'is to exceed 0.55 mIoU on damage classes, substantially narrowing the gap with state-of-the-art '
    'methods.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Second, operational sensor validation is needed to replace archival metadata with real-time sensor '
    'streams from USGS seismic networks, NOAA weather stations, and IoT flood gauges, testing the '
    'model robustness under genuine noise, packet loss, and sensor drift conditions.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Third, collecting temporally and spatially co-registered data across all three modalities would '
    'enable end-to-end validation of the fusion architecture on naturally paired samples rather than '
    'synthetic pairings.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Fourth, exploring Guided Cross Attention with synthetic caption generation for the crisis '
    'sub-model could close the gap with the current CrisisMMD state of the art while maintaining '
    'embedding quality for downstream fusion.',
    space_after=8, first_line_indent=1.27
)

add_normal_para(
    'Finally, integration with first-responder feedback through field deployment would provide '
    'operational evaluation of the explainability framework and priority classification under '
    'real disaster response conditions.',
    space_after=8, first_line_indent=1.27
)

# ============================================================
# REFERENCES (APA FORMAT)
# ============================================================
add_page_break()
add_heading_custom('References', level=0)
add_empty_lines(1)

references = [
    'Alam, F., Ofli, F., & Imran, M. (2018). CrisisMMD: Multimodal Twitter datasets from natural '
    'disasters. Proceedings of the 12th International AAAI Conference on Web and Social Media '
    '(ICWSM), 465-473.',

    'Abavisani, M., Wu, L., Hu, S., Tetreault, J., & Jaimes, A. (2020). Multimodal categorization '
    'of crisis events in social media. Proceedings of the IEEE/CVF Conference on Computer Vision '
    'and Pattern Recognition (CVPR).',

    'Allen, R. M. (2009). Real-time earthquake detection and hazard assessment by ElarmS across '
    'California. Geophysical Research Letters, 36(5).',

    'Basha, E., & Rus, D. (2008). Design of early warning flood detection systems for developing '
    'countries. Proceedings of the International Conference on Information and Communication '
    'Technologies and Development (ICTD).',

    'Benson, V., Ecker, A., & Schmid, R. (2025). Multi-scale damage assessment on xBD: Advancing '
    'building damage segmentation with Dice-Focal loss. Proceedings of the IEEE International '
    'Geoscience and Remote Sensing Symposium (IGARSS).',

    'Centre for Research on the Epidemiology of Disasters (CRED). (2024). 2023 Disasters in '
    'numbers. CRED/UCLouvain, Brussels, Belgium.',

    'Chefer, H., Gur, S., & Wolf, L. (2021). Transformer interpretability beyond attention '
    'visualization. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern '
    'Recognition (CVPR), 782-791.',

    'CORGIS Datasets Project. (2023). Earthquakes dataset. Virginia Tech. '
    'https://corgis-edu.github.io/corgis/csv/earthquakes/',

    'Dewminim Naadi. (2024). Sri Lanka Flood Risk & Inundation Dataset. Kaggle. '
    'https://www.kaggle.com/datasets/dewminimnaadi/sri-lanka-flood-risk-and-inundation-dataset',

    'Gupta, R., Hosfelt, R., Saber, S., Ashton, N., Mullan, R., Campbell, S., et al. (2019). '
    'xBD: A dataset for assessing building damage from satellite imagery. Proceedings of the '
    'IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops (CVPRW).',

    'Imran, M., Castillo, C., Diaz, F., & Vieweg, S. (2015). Processing social media messages '
    'in mass emergency: A survey. ACM Computing Surveys, 47(4), 1-38.',

    'Khaled, A. (2023). California Weather and Fire Prediction Dataset. Kaggle. '
    'https://www.kaggle.com/datasets/aleenkhaled/california-weather-and-fire-prediction-dataset',

    'Li, J., Li, D., Xiong, C., & Hoi, S. (2022). BLIP: Bootstrapping language-image pre-training '
    'for unified vision-language understanding and generation. Proceedings of the International '
    'Conference on Machine Learning (ICML).',

    'Lu, J., Batra, D., Parikh, D., & Lee, S. (2019). ViLBERT: Pretraining task-agnostic '
    'visiolinguistic representations for vision-and-language tasks. Advances in Neural Information '
    'Processing Systems (NeurIPS).',

    'Mouzannar, H., Rizk, Y., & Awad, M. (2025). Guided cross-attention for multimodal crisis '
    'classification with LLaVA-generated captions. Information Processing & Management, 62.',

    'Ofli, F., Alam, F., & Imran, M. (2020). Analysis of social media data using multimodal '
    'deep learning for disaster response. Proceedings of the 17th International Conference on '
    'Information Systems for Crisis Response and Management (ISCRAM).',

    'Olteanu, A., Castillo, C., Diaz, F., & Vieweg, S. (2014). CrisisLex: A lexicon for '
    'collecting and filtering microblogged communications in crises. Proceedings of the 8th '
    'International AAAI Conference on Weblogs and Social Media (ICWSM).',

    'Rudner, T. G. J., Russwurm, M., Fil, J., Pelich, R., Bischke, B., Kopackova, V., & '
    'Bilinski, P. (2019). Multi3Net: Segmenting flooded buildings via fusion of multiresolution, '
    'multisensor, and multitemporal satellite imagery. Proceedings of the AAAI Conference on '
    'Artificial Intelligence.',

    'Sakaki, T., Okazaki, M., & Matsuo, Y. (2010). Earthquake shakes Twitter users: Real-time '
    'event detection by social sensors. Proceedings of the 19th International Conference on '
    'World Wide Web (WWW), 851-860.',

    'Tan, H., & Bansal, M. (2019). LXMERT: Learning cross-modality encoder representations '
    'from transformers. Proceedings of the Conference on Empirical Methods in Natural Language '
    'Processing (EMNLP).',

    'The Devastator. (2023). Tropical Cyclone Tracks Dataset. Kaggle. '
    'https://www.kaggle.com/datasets/thedevastator/tropical-cyclone-tracks-dataset',

    'Tukey, J. (2023). Dataset of Earthquakes in Iran (1996-2022). Kaggle. '
    'https://www.kaggle.com/datasets/johntukey/iran-earthquake',

    'Utkarsh. (2024). NOAA Atlantic Hurricane Dataset (1975-2024). Kaggle. '
    'https://www.kaggle.com/datasets/utkarshx27/noaa-atlantic-hurricane-database',

    'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., '
    '& Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information '
    'Processing Systems (NeurIPS).',

    'Weber, E., & Kan, H. (2020). Building disaster damage assessment in satellite imagery with '
    'multi-temporal fusion. Proceedings of the IEEE International Conference on Big Data.',

    'Defense Innovation Unit. (2019). xView2 Challenge: Assessing building damage. '
    'https://xview2.org',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(f'[{i}] {ref}')
    set_font(run, size=12)

# ============================================================
# APPENDICES
# ============================================================
add_page_break()
add_heading_custom('Appendices', level=0)
add_empty_lines(1)

add_heading_custom('Appendix 1: Technology Stack', level=1)

add_normal_para('The following technologies were used in this project:', space_after=8)

tech_items = [
    'ML Framework: PyTorch >= 2.0',
    'Vision Backbone: BLIP ViT (Salesforce/blip-image-captioning-base)',
    'Text Backbone: XLM-RoBERTa (xlm-roberta-base)',
    'Segmentation: DeepLabV3+ with ResNet101',
    'API Server: FastAPI + Uvicorn',
    'XAI Explanations: GPT-4o (OpenAI API)',
    'Frontend: Vanilla JavaScript and CSS',
    'Training Hardware: Kaggle NVIDIA Tesla P100 (16 GB)',
]

for item in tech_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(f'\u2022 {item}')
    set_font(run, size=12)

add_empty_lines(1)
add_heading_custom('Appendix 2: Model Checkpoints', level=1)

add_normal_para('The trained model files and their approximate sizes:', space_after=8)

model_items = [
    'IoT Classifier (IOT/models/iot_model.pth): ~180 KB, contains AdaptiveIoTClassifier state_dict, config, and validation metrics',
    'Crisis Classifier (crisis/best_adaptive_model.pth): ~450 MB, contains AdaptiveFusionClassifier state_dict including frozen BLIP and XLM-RoBERTa weights',
    'Fusion Layer (fusion/fusion_model.pth): ~2 MB, contains FusionLayer state_dict',
]

for item in model_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(f'\u2022 {item}')
    set_font(run, size=12)

add_empty_lines(1)
add_heading_custom('Appendix 3: Directory Structure', level=1)

dir_structure = """Project Root/
    IOT/
        train_iot.py              - IoT model training script
        models/iot_model.pth      - Trained IoT model checkpoint
        datasets/                 - Processed CSV files from 6 Kaggle datasets
    crisis/
        train_crisis_code.ipynb   - Crisis model training notebook
        server.py                 - AdaptiveFusionClassifier definition
        best_adaptive_model.pth   - Trained crisis model checkpoint
    XBD/
        final-xbd-deeplabv3plus.ipynb - xBD model training notebook
    fusion/
        fusion_layer.py           - TriFusionLayer and FusionLayer definitions
        train_fusion.py           - Fusion layer training script
        pipeline.py               - Inference pipeline
        xai.py                    - Explainability module
        iot_predictor.py          - IoT inference wrapper
        fusion_model.pth          - Trained fusion model checkpoint
    outputs-paper/                - Training outputs and metrics
    paper/                        - IEEE conference paper versions"""

p = doc.add_paragraph()
run = p.add_run(dir_structure)
set_font(run, size=10, name='Courier New')

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'thesis', 'VIT_AP_Thesis_TriModal_Disaster_Intelligence.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Thesis saved to: {output_path}')
print(f'Total sections: Cover, Title, Declaration, Certificate, Abstract, Acknowledgement,')
print(f'  Table of Contents, List of Tables, List of Figures, List of Abbreviations,')
print(f'  Chapter 1-6, References, Appendices')
